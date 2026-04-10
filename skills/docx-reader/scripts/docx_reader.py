#!/usr/bin/env python3
"""
docx-reader — extract the text content of a Word (.docx) file.

One of the four readers in the `document-readers` plugin. Companion
to the `document-search` skill: use document-search to locate the
file, then this skill to read its content.

Usage:
    docx_reader.py <file.docx> [--max-bytes N]
                               [--format text|json]
                               [--metadata-only]

Run with --help for the full option list.
"""
from __future__ import annotations

import argparse
import json
import sys
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Optional

sys.path.insert(0, str(Path(__file__).resolve().parent))
import _preflight  # noqa: E402


@dataclass
class ReadResult:
    path: str
    extension: str = ".docx"
    kind: str = "docx"
    size: int = 0
    content: Optional[str] = None
    truncated: bool = False
    metadata: dict = field(default_factory=dict)
    error: Optional[str] = None
    missing_imports: list[str] = field(default_factory=list)
    install_guide: Optional[str] = None


class MissingDependency(RuntimeError):
    def __init__(self, import_name: str, feature: str):
        self.import_name = import_name
        self.feature = feature
        super().__init__(f"missing optional package '{import_name}' for {feature}")


def read_docx(path: Path) -> tuple[str, dict]:
    try:
        import docx  # type: ignore  # python-docx
    except ImportError as e:
        raise MissingDependency("docx", ".docx reading") from e
    document = docx.Document(str(path))
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    metadata = {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
    }
    return "\n".join(parts), metadata


def read_file(
    path: Path,
    *,
    max_bytes: int,
    metadata_only: bool,
) -> ReadResult:
    result = ReadResult(path=str(path), size=path.stat().st_size)
    try:
        content, meta = read_docx(path)
        result.metadata = meta
    except MissingDependency as e:
        result.error = str(e)
        result.missing_imports = [e.import_name]
        result.install_guide = _preflight.format_install_guide(
            [e.import_name],
            feature="reading .docx files",
            skill_name="docx-reader",
        )
        return result
    except Exception as e:  # noqa: BLE001
        result.error = f"read failed: {e}"
        return result

    if content is not None and len(content.encode("utf-8", errors="ignore")) > max_bytes:
        content = content.encode("utf-8", errors="ignore")[:max_bytes].decode("utf-8", errors="replace")
        result.truncated = True

    if not metadata_only:
        result.content = content
    return result


def render_text(r: ReadResult) -> str:
    out: list[str] = []
    out.append("# docx-reader")
    out.append(f"path:      {r.path}")
    out.append(f"size:      {r.size} bytes")
    if r.truncated:
        out.append("truncated: yes (output capped by --max-bytes)")
    if r.metadata:
        out.append("metadata:")
        for k, v in r.metadata.items():
            out.append(f"  {k}: {v}")
    out.append("")

    if r.error:
        out.append(f"error: {r.error}")
        if r.install_guide:
            out.append("")
            out.append("## Install guide")
            out.append(r.install_guide)
        return "\n".join(out)

    if r.content is None:
        out.append("(metadata-only mode; no content returned)")
        return "\n".join(out)

    out.append("## Content")
    out.append(r.content)
    return "\n".join(out)


def render_json(r: ReadResult) -> str:
    return json.dumps(asdict(r), ensure_ascii=False, indent=2)


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="docx_reader.py",
        description="Extract the text content of a Word (.docx) file, "
                    "including paragraphs and table cells.",
    )
    p.add_argument("file", help="Path to the .docx file to read.")
    p.add_argument("--max-bytes", type=int, default=200_000,
                   help="Maximum bytes of content returned (default: 200000).")
    p.add_argument("--format", choices=("text", "json"), default="text",
                   help="Output format (default: text).")
    p.add_argument("--metadata-only", action="store_true",
                   help="Return only metadata (paragraph + table counts) without the content.")
    return p


def main(argv: Optional[list[str]] = None) -> int:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")  # type: ignore[attr-defined]
        except (AttributeError, ValueError):
            pass

    args = build_parser().parse_args(argv)
    path = Path(args.file)
    if not path.exists():
        print(f"error: file does not exist: {path}", file=sys.stderr)
        return 2
    if not path.is_file():
        print(f"error: not a regular file: {path}", file=sys.stderr)
        return 2
    if path.suffix.lower() != ".docx":
        print(f"warning: file does not have a .docx extension: {path}", file=sys.stderr)

    result = read_file(
        path,
        max_bytes=args.max_bytes,
        metadata_only=args.metadata_only,
    )
    if args.format == "json":
        print(render_json(result))
    else:
        print(render_text(result))
    return 0 if result.error is None else 2


if __name__ == "__main__":
    sys.exit(main())
