#!/usr/bin/env python3
"""
docx-inspector — structural inspection of a Word (.docx) file.

Unlike docx-reader, which returns *text*, this skill returns *structure* —
the deterministic facts an LLM cannot reliably infer from extracted
text:

  * Core document properties (title, author, dates, revision, ...)
    plus computed counts (paragraphs, words, tables, sections)
  * Tracked changes — every w:ins / w:del XML element, with author,
    date, and the actual inserted or deleted text. python-docx does
    not expose these through its high-level API, so we walk the lxml
    tree directly.
  * Heading hierarchy outline, with skipped-level detection (flags
    H1→H3 jumps, documents that don't start at H1, etc.)

Sibling skills in the `document-inspector` plugin handle other formats
(pdf-inspector for .pdf, xlsx-inspector for .xlsx). This skill is
.docx-only by design.

Usage:
    docx_inspector.py <file> --feature metadata|changes|structure|all
                             [--format text|json]

Run with --help for the full option list.
"""
from __future__ import annotations

import argparse
import json
import sys
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Any, Optional

# Self-contained install-guide helper.
sys.path.insert(0, str(Path(__file__).resolve().parent))
import _preflight  # noqa: E402


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class InspectResult:
    path: str
    extension: str = ".docx"
    kind: str = "docx"
    feature: str = "all"
    size: int = 0
    metadata: dict = field(default_factory=dict)
    changes: dict = field(default_factory=dict)
    structure: dict = field(default_factory=dict)
    error: Optional[str] = None
    missing_imports: list[str] = field(default_factory=list)
    install_guide: Optional[str] = None


class MissingDependency(RuntimeError):
    def __init__(self, import_name: str, feature: str):
        self.import_name = import_name
        self.feature = feature
        super().__init__(f"missing optional package '{import_name}' for {feature}")


# ---------------------------------------------------------------------------
# Open + dispatch
# ---------------------------------------------------------------------------

FEATURES = ("metadata", "changes", "structure", "all")

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP = {"w": WORD_NS}


def _open_docx(path: Path):
    try:
        import docx  # type: ignore
    except ImportError as e:
        raise MissingDependency("docx", "docx inspection") from e
    return docx.Document(str(path))


def _w(name: str) -> str:
    """Fully-qualified w: attribute name for lxml."""
    return f"{{{WORD_NS}}}{name}"


def _datetime_to_iso(value: Any) -> Optional[str]:
    if value is None:
        return None
    iso = getattr(value, "isoformat", None)
    if callable(iso):
        try:
            return iso()
        except Exception:  # noqa: BLE001
            pass
    return str(value)


# ---------------------------------------------------------------------------
# Feature: metadata
# ---------------------------------------------------------------------------

_CORE_PROPS = (
    "title", "author", "subject", "keywords", "comments",
    "category", "content_status",
    "last_modified_by", "revision", "version", "language",
    "identifier",
)


def _docx_metadata(doc) -> dict:
    """
    Extract core properties from doc.core_properties plus computed
    structural counts. Fields present in every output (never silently
    dropped), missing values become None.
    """
    out: dict[str, Any] = {}
    props = doc.core_properties

    for name in _CORE_PROPS:
        value = getattr(props, name, None)
        out[name] = None if value in (None, "") else str(value)

    out["created"] = _datetime_to_iso(getattr(props, "created", None))
    out["modified"] = _datetime_to_iso(getattr(props, "modified", None))
    out["last_printed"] = _datetime_to_iso(getattr(props, "last_printed", None))

    # Computed structural facts.
    paragraph_count = len(doc.paragraphs)
    word_count = 0
    for p in doc.paragraphs:
        word_count += len(p.text.split())
    out["paragraph_count"] = paragraph_count
    out["word_count"] = word_count
    out["table_count"] = len(doc.tables)
    try:
        out["section_count"] = len(doc.sections)
    except Exception:  # noqa: BLE001
        out["section_count"] = None
    return out


# ---------------------------------------------------------------------------
# Feature: changes (tracked changes via raw XML walk)
# ---------------------------------------------------------------------------

_TEXT_CAP = 200
_CHANGES_CAP = 200


def _gather_text(elem, tag_local: str) -> str:
    """
    Concatenate the .text of every descendant matching {w:ns}tag_local.
    """
    parts: list[str] = []
    for t in elem.findall(f".//w:{tag_local}", NSMAP):
        if t.text:
            parts.append(t.text)
    return "".join(parts)


def _truncate(s: str, cap: int = _TEXT_CAP) -> str:
    if len(s) <= cap:
        return s
    return s[: cap - 1] + "…"


def _docx_changes(doc) -> dict:
    """
    Walk the document XML for w:ins and w:del elements (tracked
    insertions and deletions). python-docx does not expose these
    through its high-level API — we go straight to the underlying
    lxml tree.

    Returns:
      {
        "has_tracked_changes": bool,
        "insertion_count": int,
        "deletion_count": int,
        "authors": [sorted unique authors],
        "changes": [ {type, author, date, text, id}, ... ],  # capped
        "truncated": bool,
      }
    """
    root = doc.element  # lxml Element for the w:document
    inserts = list(root.findall(".//w:ins", NSMAP))
    deletes = list(root.findall(".//w:del", NSMAP))

    authors: set[str] = set()
    collected: list[dict[str, Any]] = []

    def _record(change_type: str, el) -> None:
        author = el.get(_w("author"))
        date = el.get(_w("date"))
        change_id = el.get(_w("id"))
        if change_type == "insert":
            text = _gather_text(el, "t")
        else:
            text = _gather_text(el, "delText")
        if author:
            authors.add(author)
        collected.append({
            "type": change_type,
            "author": author,
            "date": date,
            "text": _truncate(text) if text else "",
            "id": change_id,
        })

    for el in inserts:
        _record("insert", el)
    for el in deletes:
        _record("delete", el)

    total = len(collected)
    truncated = total > _CHANGES_CAP
    if truncated:
        collected = collected[:_CHANGES_CAP]

    return {
        "has_tracked_changes": bool(inserts or deletes),
        "insertion_count": len(inserts),
        "deletion_count": len(deletes),
        "authors": sorted(authors),
        "changes": collected,
        "truncated": truncated,
    }


# ---------------------------------------------------------------------------
# Feature: structure (heading outline + hierarchy checks)
# ---------------------------------------------------------------------------

_OUTLINE_CAP = 200


def _parse_heading_level(style_name: str) -> Optional[int]:
    """
    Return the integer level of a 'Heading N' style, or None for any
    other style. Accepts things like 'Heading 1' and 'heading 2'.
    """
    if not style_name:
        return None
    lowered = style_name.lower()
    if not lowered.startswith("heading "):
        return None
    try:
        return int(lowered.split()[-1])
    except (ValueError, IndexError):
        return None


def _docx_structure(doc) -> dict:
    """
    Build a heading outline and flag hierarchy issues (skipped levels,
    first-heading-not-H1). Does not fail on documents with no headings.
    """
    outline: list[dict[str, Any]] = []
    issues: list[dict[str, Any]] = []

    previous_level: Optional[int] = None
    max_depth = 0
    first_heading_seen = False

    for idx, paragraph in enumerate(doc.paragraphs):
        style_name = ""
        try:
            style_name = paragraph.style.name
        except Exception:  # noqa: BLE001
            pass
        level = _parse_heading_level(style_name)
        if level is None:
            continue

        text = paragraph.text or ""
        if level > max_depth:
            max_depth = level

        if not first_heading_seen:
            first_heading_seen = True
            if level != 1:
                issues.append({
                    "index": idx,
                    "text": _truncate(text, 120),
                    "level": level,
                    "previous_level": None,
                    "issue": f"document does not start at H1 (first heading is H{level})",
                })
        else:
            if previous_level is not None and level > previous_level + 1:
                issues.append({
                    "index": idx,
                    "text": _truncate(text, 120),
                    "level": level,
                    "previous_level": previous_level,
                    "issue": (
                        f"skipped from H{previous_level} to H{level} "
                        f"(missing H{previous_level + 1})"
                    ),
                })

        if len(outline) < _OUTLINE_CAP:
            outline.append({
                "level": level,
                "text": _truncate(text, 120),
                "index": idx,
            })

        previous_level = level

    return {
        "heading_count": sum(1 for _ in (
            p for p in doc.paragraphs
            if _parse_heading_level(getattr(getattr(p, "style", None), "name", "")) is not None
        )),
        "max_depth": max_depth,
        "outline": outline,
        "outline_truncated": len(outline) >= _OUTLINE_CAP,
        "hierarchy_issues": issues,
    }


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

def inspect(path: Path, feature: str) -> InspectResult:
    result = InspectResult(
        path=str(path),
        extension=path.suffix.lower(),
        kind="docx" if path.suffix.lower() == ".docx" else "unknown",
        feature=feature,
        size=path.stat().st_size,
    )

    if result.extension != ".docx":
        result.error = (
            f"docx-inspector only accepts .docx files "
            f"(got {result.extension or '(none)'}) — use pdf-inspector "
            f"or xlsx-inspector for other formats"
        )
        return result

    try:
        doc = _open_docx(path)
    except MissingDependency as e:
        result.error = str(e)
        result.missing_imports = [e.import_name]
        result.install_guide = _preflight.format_install_guide(
            [e.import_name],
            feature="inspecting .docx files",
            skill_name="docx-inspector",
        )
        return result
    except Exception as e:  # noqa: BLE001
        result.error = f"failed to open DOCX: {e}"
        return result

    try:
        if feature in ("metadata", "all"):
            result.metadata = _docx_metadata(doc)
        if feature in ("changes", "all"):
            result.changes = _docx_changes(doc)
        if feature in ("structure", "all"):
            result.structure = _docx_structure(doc)
    except Exception as e:  # noqa: BLE001
        result.error = f"inspection failed: {e}"
        return result

    return result


# ---------------------------------------------------------------------------
# Rendering
# ---------------------------------------------------------------------------

def render_text(r: InspectResult) -> str:
    out: list[str] = []
    out.append("# docx-inspector")
    out.append(f"path:      {r.path}")
    out.append(f"size:      {r.size} bytes")
    out.append(f"kind:      {r.kind}")
    out.append(f"feature:   {r.feature}")
    out.append("")

    if r.error:
        out.append(f"error: {r.error}")
        if r.install_guide:
            out.append("")
            out.append("## Install guide")
            out.append(r.install_guide)
        return "\n".join(out)

    if r.feature in ("metadata", "all") and r.metadata:
        out.append("## Metadata")
        ordered_keys = (
            "title", "author", "subject", "keywords", "comments",
            "category", "content_status",
            "created", "modified", "last_printed",
            "last_modified_by", "revision", "version", "language",
            "identifier",
            "paragraph_count", "word_count", "table_count", "section_count",
        )
        for key in ordered_keys:
            if key in r.metadata:
                val = r.metadata[key]
                out.append(f"  {key}: {val if val is not None else '(none)'}")
        out.append("")

    if r.feature in ("changes", "all") and r.changes:
        out.append("## Tracked changes")
        if not r.changes.get("has_tracked_changes"):
            out.append("  (no tracked changes)")
        else:
            out.append(f"  insertion_count: {r.changes.get('insertion_count', 0)}")
            out.append(f"  deletion_count:  {r.changes.get('deletion_count', 0)}")
            authors = r.changes.get("authors") or []
            if authors:
                out.append(f"  authors: {', '.join(authors)}")
            if r.changes.get("truncated"):
                total = r.changes.get("insertion_count", 0) + r.changes.get("deletion_count", 0)
                out.append(f"  (showing first {_CHANGES_CAP} of {total} changes)")
            out.append("")
            out.append("  changes:")
            for c in r.changes.get("changes", []):
                author = c.get("author") or "(unknown)"
                date = c.get("date") or "(no date)"
                text = c.get("text") or "(empty)"
                kind = c.get("type", "?")
                out.append(f"    [{kind}] {author} @ {date}: {text!r}")
        out.append("")

    if r.feature in ("structure", "all") and r.structure:
        out.append("## Structure")
        out.append(f"  heading_count: {r.structure.get('heading_count', 0)}")
        out.append(f"  max_depth:     {r.structure.get('max_depth', 0)}")
        issues = r.structure.get("hierarchy_issues") or []
        if issues:
            out.append(f"  hierarchy_issues: {len(issues)}")
            for iss in issues[:20]:
                out.append(f"    - index {iss.get('index')}: {iss.get('issue')}")
                out.append(f"      text: {iss.get('text')!r}")
        outline = r.structure.get("outline") or []
        if outline:
            out.append("")
            out.append("  outline:")
            for h in outline[:80]:
                indent = "  " * h.get("level", 1)
                out.append(f"    {indent}H{h.get('level')}: {h.get('text')}")
            if len(outline) > 80:
                out.append(f"    ... ({len(outline) - 80} more headings)")
        out.append("")

    return "\n".join(out).rstrip() + "\n"


def render_json(r: InspectResult) -> str:
    return json.dumps(asdict(r), ensure_ascii=False, indent=2)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="docx_inspector.py",
        description="Structural inspection of a single Word (.docx) file — "
                    "core properties, tracked changes, and heading-hierarchy "
                    "outline. Returns deterministic facts an LLM cannot derive "
                    "from extracted text.",
    )
    p.add_argument("file", help="Path to the .docx file to inspect.")
    p.add_argument("--feature", choices=FEATURES, default="all",
                   help="Which inspection to run (default: all).")
    p.add_argument("--format", choices=("text", "json"), default="text",
                   help="Output format (default: text).")
    return p


def main(argv: Optional[list[str]] = None) -> int:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")  # type: ignore[attr-defined]
        except (AttributeError, ValueError):
            pass

    args = build_parser().parse_args(argv)
    path = Path(args.file)
    if not path.exists():
        print(f"error: file does not exist: {path}", file=sys.stderr)
        return 2
    if not path.is_file():
        print(f"error: not a regular file: {path}", file=sys.stderr)
        return 2

    result = inspect(path, feature=args.feature)

    if args.format == "json":
        print(render_json(result))
    else:
        print(render_text(result))
    return 0 if result.error is None else 2


if __name__ == "__main__":
    sys.exit(main())
